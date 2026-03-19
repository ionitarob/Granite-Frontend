from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_endpoint(doc, method, url, desc):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f"[{method}] {url}")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)
    p.add_run(f": {desc}")

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('Documentación Técnica Avanzada: Granite-Frontend', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Mapeo exhaustivo de pantallas, funcionalidades y consumo de APIs.', style='Subtitle')

    # Seccion 1
    doc.add_heading('1. Introducción', level=1)
    doc.add_paragraph(
        "Este documento técnico expone en profundidad todos los módulos internos de la aplicación Flutter Granite (Frontend), "
        "las pantallas que componen sus procesos, los flujos del operario en cada una y los Endpoints exactos que "
        "interactúan con Granite-Backend. Esta documentación sirve como referencia para administradores, arquitectos de software y auditores del sistema."
    )

    # Seccion 2: OrderOps
    doc.add_heading('2. Módulo OrderOps (Operaciones de Órdenes)', level=1)
    doc.add_paragraph('Archivos y lógica base: lib/screens/orderops/... y services/orderops_service.dart')
    doc.add_paragraph(
        "Flujo y qué se hace en las pantallas: En este módulo, el usuario visualiza las órdenes emitidas al almacén "
        "o a la estación de trabajo. El operario hace clic sobre una orden en la lista, entrando en la pantalla 'OrderDetailScreen' "
        "(Detalle de Orden). Aquí puede validar que los SKUs y las cantidades coincidan con el pedido físico. "
        "También puede arrancar el motor de decisión local 'Triage' para categorizar un pedido por 'Familias'. Si la familia "
        "detectada requiere captura de números de serie, la propia interfaz navega e integra embebido el módulo Serials (SerialLinkScreen). "
        "Por la pantalla de detalle se pueden documentar discrepancias ('Observaciones') y adjuntar fotos ('Archivos Evidencia')."
    )
    doc.add_heading('Endpoints Principales Consumidos:', level=2)
    add_endpoint(doc, 'GET', '/orderops/agent-orders?limit=X', 'Recupera el listado maestro de órdenes asignables/visibles al turno, para llenar el Grid inicial del almacén.')
    add_endpoint(doc, 'GET', '/orderops/agent-orders/{idnbr}', 'Extrae la vista 360 de un pedido (cabecera, ítems, servicios asociados y estado en tiempo real).')
    add_endpoint(doc, 'PATCH', '/orderops/agent-orders/{idnbr}/status', 'Cambia el ciclo de vida del pedido (Ej. "Pending" a "Executing" o a "Finished"). Se dispara por botones en la UI superior.')
    add_endpoint(doc, 'POST', '/orderops/triage/run', 'Envía parte del pedido para que el backend evalúe reglas de negocio (ej. SKU "MANIP-01" -> Cae en Familia "Manipulación y Etiquetado").')
    add_endpoint(doc, 'GET', '/orderops/agent-orders/{id}/photos', 'Trae metadatos de las imágenes adjuntas subidas como Evidencia del pedido.')
    add_endpoint(doc, 'MULTIPART-POST', '/orderops/agent-orders/{id}/photos', 'Sube un archivo binario capturado por cámara web o disco duro para asociarlo fidedignamente al pedido físico.')
    add_endpoint(doc, 'GET/POST/DELETE', '/orderops/observations/*', 'Buzón bidireccional de incidencias e instrucciones operativas (ej. "Caja venía golpeada").')

    # Seccion 3: Serials
    doc.add_heading('3. Módulo de Serials (Gestión Avanzada de Series)', level=1)
    doc.add_paragraph('Archivos y lógica base: lib/screens/serials/... y services/mask_service.dart')
    doc.add_paragraph(
        "Flujo y qué se hace en las pantallas: Este es el motor logístico principal de trazabilidad de ítems individuales. "
        "Cuando el operario abre la pantalla 'SerialLinkScreen' (dentro del detalle o independientemente), escanea secuencialmente: "
        "1. El Número de Orden (o se pasa autocompletado). 2. El Código de Inventario temporal libre. 3. El Serial real del dispositivo. "
        "El módulo cuenta con controles defensivos antierrores: escanea el texto rápidamente, y la aplicación llama a validadores de MÁSCARAS en el backend "
        "(para alertar si el texto escaneado resulta ser un part number en lugar de un MAC o SN auténtico). "
        "Otra pantalla crítica es 'SerialChange' (Cambio de Serial), donde se documenta la sustitución de un componente defectuoso por uno en buen estado, reasignando su inventario."
    )
    doc.add_heading('Endpoints Principales Consumidos:', level=2)
    add_endpoint(doc, 'POST', '/serials/order-info', 'Calcula cuántas unidades se declararon en la orden y devuelve un JSON al front-end para renderizar la tabla o "slots" de escaneo en blanco pendientes.')
    add_endpoint(doc, 'GET', '/serials/next-available', 'Técnica de agilidad operativa; precarga el código de inventario libre siguiente mientras el operario manipula la pistola de escaneo.')
    add_endpoint(doc, 'POST', '/serials/masks/check', 'Valida el texto escaneado contra expresiones regulares/máscaras prohibitivas registradas en backend. Bloquea si detecta formato erróneo.')
    add_endpoint(doc, 'POST', '/serials/match', 'El "cierre" de la fila. Inserta definitivamente en la base de datos (tabla serial_to_orden) la vinculación Número de Orden -> Serial Registrado.')
    add_endpoint(doc, 'POST', '/serials/finish-order-upload', 'Llamada clave de fin de trabajo. Al clicar "Finalizar Orden", obliga al backend a tomar el acumulado, generar un documento Excel nativo (librería openpyxl), formatearlo (ej. num_orden, serial, inventory_code), empaquetarlo y enviarlo por SFTP corporativo para los sistemas externos.')
    add_endpoint(doc, 'GET', '/serials/matches/export?num_orden={X}', 'Llamada usada al dar clic en "Descargar/Exportar Vista" en el fronted. Descarga el Excel volátil al navegado web/app local. Usa internamente un SQL REPLACE(num_orden, "-", "") para máxima resiliencia ante errores de sintaxis u homogeneización (ej. 291362511 vs 29-13625-11).')
    add_endpoint(doc, 'GET', '/serials/export-serial-changes?nr_orden={X}', 'Exportación específica pero aplicada exclusivamente para la tabla de trazabilidad médica de piezas llamada SerialChange.')

    # Seccion 4: Amazon
    doc.add_heading('4. Módulo Logístico: Amazon', level=1)
    doc.add_paragraph('Archivos y lógica base: lib/screens/amazon/...')
    doc.add_paragraph(
        "Flujo y qué se hace en las pantallas: Se compone en subsecciones de Índice de Calidad y Control de Inventario ("
        "Inventory Control y Grading Screen). En Amazon Grading, el evaluador introduce lote y escanea la caja Amazon. "
        "Aparece una lista de revisión técnica; dependiendo de los defectos seleccionados, la tablet le asigna automáticamente un Grado (A, B, C). "
        "En Amazon Inventory Control, los encajadores componen (paletizan) sub-cajas en WPL/Pallets contenedores validando pesos, agencias y destinos logísticos FBA."
    )
    doc.add_heading('Endpoints Principales Consumidos:', level=2)
    add_endpoint(doc, 'GET', '/amz/inventory/wpl/{wplId}/export', 'Extrae los manifiestos ad-hoc del conteo de referencias alojadas de un pallet en concreto para control de embarque o aduanas.')
    add_endpoint(doc, 'GET', '/amz/inventory/export', 'Extracción genérica masiva y sumarizada de inventario registrado bajo los flujos del ecosistema Amazon dentro de planta.')
    
    # Seccion 5: RRHH 
    doc.add_heading('5. Módulo de Recursos Humanos y Tiempos (Job Selector)', level=1)
    doc.add_paragraph('Archivos y lógica base: lib/screens/rrhh/... ')
    doc.add_paragraph(
        "Flujo y qué se hace en las pantallas: Todo usuario operativo interactúa de inmediato con esta pantalla al iniciar turno. "
        "Se escanea el ID de Puesto / Estación de Trabajo, validando la geolocalización o sector (ej. Recepción Docks, Triaje Servidores, Empaquetado Amazon). "
        "Esto levanta un temporizador silencioso en el backend. Cuando el operario se desplaza a otro proyecto y ficha de nuevo, el reloj detiene su cronómetro "
        "anterior, generando estadísticas financieras de rendimiento (Performance OEE)."
    )
    doc.add_heading('Endpoints Principales Consumidos:', level=2)
    add_endpoint(doc, 'GET', '/jobs/export', 'Punto de recolección exportable tabular (CSV/Excel) de rendimientos y marcajes ("punzonamientos") a lo largo del tiempo, facilitando auditorías al departamento de HR y Operaciones.')

    # Seccion 6: Modulos Secundarios
    doc.add_heading('6. Otros Módulos Departamentales (Sentinels, Previs, Servidores)', level=1)
    doc.add_paragraph('Archivos base: previ_service.dart, server_registro_service.dart, igualdad_api.dart')
    doc.add_paragraph(
        "Previs: Mantenimiento y prevención logística; pantallas que usan GET/PATCH/DELETE básicos contra /previs. "
        "Igualdad: Gestión de expediciones para importación documental a clientes; cuenta con llamadas agresivas de subida masiva: POST /igualdad/entrada/importar_registro y su respectiva exportable /igualdad/exportar_expedicion. "
        "Servers: Flujo simplificado pero robusto mediante POST para inyección de perfiles de borrado seguro en máquinas físicas y reportes de QA."
    )

    doc.add_heading('Resumen Funcional y Flujos Arquitecturales', level=1)
    doc.add_paragraph(
        "A nivel de desarrollo, todas estas pantallas residen en un Single Page Application impulsada por Flutter Web/Desktop, "
        "que se actualiza síncronamente mediante Provider (API y Estado Global). El conector base, ApiClient, es el encargado de adosar el "
        "token JWT en cada 'Authorization: Bearer' inyectado dinámicamente. Esto garantiza que todos los endpoints mencionados actúen "
        "comprobando la jerarquía de roles, aislando o habilitando acciones destructivas según la acreditación del usuario corporativo."
    )

    doc.save('/Users/rionita_work/Development/GRANITE/Granite-Frontend/Documentacion_Avanzada_Granite.docx')
    print("Advanced Document successfully created!")

if __name__ == "__main__":
    main()
