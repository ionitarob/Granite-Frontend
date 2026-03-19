from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('Documentación de la Aplicación Granite', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Generado automáticamente. Sistema de Control Operativo y Calidad.', style='Subtitle')

    # Seccion 1
    doc.add_heading('1. Introducción', level=1)
    doc.add_paragraph(
        "Granite es una plataforma integral diseñada para la gestión, operaciones y control de calidad "
        "de diversas líneas de negocio tecnológicas y logísticas. Está dividida en múltiples módulos que "
        "cubren desde operaciones de órdenes (OrderOps) hasta la gestión de inventario, auditoría de calidad, "
        "recursos humanos y asistencia interactiva mediante Inteligencia Artificial."
    )

    # Seccion 2
    doc.add_heading('2. Interfaz Principal y Navegación', level=1)
    p2_1 = doc.add_paragraph()
    p2_1.add_run('Dashboard (Panel de Control): ').bold = True
    p2_1.add_run('La pantalla principal ofrece un resumen de métricas clave operativas (producción, horas de actividad, eficiencia) y acceso rápido a los diferentes módulos según el rol del usuario (Operario, Recursos Humanos, Administrador).')
    
    p2_2 = doc.add_paragraph()
    p2_2.add_run('Sidebar (Barra Lateral): ').bold = True
    p2_2.add_run('Permite la navegación rápida entre los distintos módulos: Dashboard, OrderOps, Amazon, RRHH, Igualdad, Serials, Xiaomi, etc.')
    
    p2_3 = doc.add_paragraph()
    p2_3.add_run('Asistente IA (AIsphere/ChatPanel): ').bold = True
    p2_3.add_run('Un asistente inteligente integrado y flotante en la interfaz que permite a los usuarios hacer preguntas sobre la documentación técnica y los procedimientos operativos internos.')

    # Seccion 3
    doc.add_heading('3. Módulos y Funcionalidades', level=1)

    doc.add_heading('3.1 OrderOps (Operaciones de Pedidos)', level=2)
    doc.add_paragraph('Gestión completa del ciclo de vida de los pedidos que procesa el almacén.')
    doc.add_paragraph('Registro de Órdenes: Visualización del detalle de cada pedido, incluyendo número de orden estructurado, cliente, SKUs a procesar, cantidades y datos de viabilidad (costes/márgenes).', style='List Bullet')
    doc.add_paragraph('Cambio de Estado: Transición manual y automática de pedidos a través de diferentes fases (Validada, Pendiente, En Ejecución, Parada, Finalizada, Facturada).', style='List Bullet')
    doc.add_paragraph('Vinculación a Familias: Categorización de órdenes por procesos (ej. Ordenadores Servidores, Manipulación y Etiquetado, Cambio de Serial) para desencadenar distintos flujos lógicos.', style='List Bullet')
    doc.add_paragraph('Archivos de Evidencia: Subida y visualización de archivos adjuntos en distintos formatos y fotos para registro de calidad.', style='List Bullet')

    doc.add_heading('3.2 Gestión de Seriales (Serials)', level=2)
    doc.add_paragraph('Vinculación (Match): Escaneo masivo de números de serie y asociación a códigos de inventario (Inventory Code) con validación avanzada contra máscaras de fábrica preventivas.', style='List Bullet')
    doc.add_paragraph('Cambio de Serial: Herramienta dedicada a la sustitución física y lógica de números de serie averiados por unos nuevos, enlazando la caja y evidencia visual o técnica.', style='List Bullet')
    doc.add_paragraph('Historial: Explorador para seguir la trazabilidad de cualquier MAC o equipo a lo largo del tiempo.', style='List Bullet')
    doc.add_paragraph('Exportación Automatizada: Los reportes de serie finales se consolidan en archivos Excel .xlsx y se suben al servidor interno SFTP cuando la orden termina con éxito.', style='List Bullet')

    doc.add_heading('3.3 Módulo Amazon', level=2)
    doc.add_paragraph('Configuraciones especiales y estándares de calidad para procesamiento logístico de envíos (FBA/Vendor).')
    doc.add_paragraph('Auditoría y Grading (Quality Index): Interfaz de peritaje del estado físico y de empaquetado del stock (ej. Grados estéticos A, B, C, Daños menores).', style='List Bullet')
    doc.add_paragraph('Control de Inventario (Inventory Control): Rastreo integral por matrícula Pallet / WPL, agencias y cajas maestras.', style='List Bullet')

    doc.add_heading('3.4 Recursos Humanos (RRHH)', level=2)
    doc.add_paragraph('Módulo de fichaje, tiempos y control de operarios en estación.')
    doc.add_paragraph('Control de Turnos y Puesto (Job Selector): Cada empleado reporta llegada y qué tareas va a desempeñar activamente.', style='List Bullet')
    doc.add_paragraph('Métricas de Productividad: Se calculan rendimientos basados en lo extraído de las lecturas y su tiempo imputado.', style='List Bullet')

    doc.add_heading('3.5 Otros Módulos Desplegables', level=2)
    doc.add_paragraph('Analisis y Serveis: Gestión de reportes de transacciones de servicio extra.', style='List Bullet')
    doc.add_paragraph('Igualdad: Control y empaquetado de expediciones al gobierno y subcontratas.', style='List Bullet')
    doc.add_paragraph('Sentinel for Imaging: Configuración y flujos para herramientas de clonación y flasheo de terminales.', style='List Bullet')
    doc.add_paragraph('Servidores: Secuencia especializada para máquinas grandes, pruebas de "Burn-in" y borrado seguro y certificado ("Wipe").', style='List Bullet')
    doc.add_paragraph('Xiaomi: Etiquetador especial simplificado para terminales de dicha firma.', style='List Bullet')

    # Seccion 4
    doc.add_heading('4. Arquitectura y Seguridad', level=1)
    doc.add_paragraph(
        "Granite-Frontend está desarrollado íntegramente en Flutter (Dart), brindando agilidad nativa y una interfaz unificada adaptativa (Responsive) a dispositivos móviles comerciales de la empresa y monitores industriales de almacén."
    )
    doc.add_paragraph(
        "La interacción se lleva a cabo contra el ecosistema Granite-Backend a través del protocolo HTTPS y en base JWT. Todo acceso está jerarquizado por roles en el ecosistema:"
    )
    p4 = doc.add_paragraph()
    p4.add_run('Operario Básico: ').bold = True
    p4.add_run('Visibilidad estrecha orientada al trabajo individual.')
    p42 = doc.add_paragraph()
    p42.add_run('Operario Avanzado: ').bold = True
    p42.add_run('Capacidad gestora de cajas u órdenes genéricas locales.')
    p43 = doc.add_paragraph()
    p43.add_run('Administrador: ').bold = True
    p43.add_run('Apertura de métricas financieras de viabilidad, auditorías totales y reversión de errores globales en bases de datos internas.')

    doc.save('/Users/rionita_work/Development/GRANITE/Granite-Frontend/Documentacion_Granite_App.docx')
    print("Document successfully created!")

if __name__ == "__main__":
    main()
